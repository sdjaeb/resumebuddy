from docx import Document
from docx.shared import Pt
import os

class DocxGenerator:
    @staticmethod
    def generate_resume(text: str, output_path: str):
        """
        Generates a reasonably formatted .docx resume from plain text.
        Assumes the text is structured with clear headings.
        """
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
                
            # Simple heuristic for headings (all caps or short lines with no ending punctuation)
            if (line.isupper() and len(line) < 50) or (line.startswith('##') or line.startswith('#')):
                # Clean up markdown headers if present
                clean_line = line.lstrip('#').strip()
                p = doc.add_paragraph()
                run = p.add_run(clean_line)
                run.bold = True
                run.font.size = Pt(14)
            else:
                doc.add_paragraph(line)

        doc.save(output_path)
        return output_path
