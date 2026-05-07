from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_resume_docx(resume_text, filename):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = resume_text.split('\n')
    
    # Header
    name = doc.add_paragraph("STEPHEN JAEB")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    contact = doc.add_paragraph("Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(0)
    
    links = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_after = Pt(12)

    current_section = None
    
    for line in lines[4:]: # Skip name and contact
        line = line.strip()
        if not line:
            continue
            
        if line in ["PROFESSIONAL SUMMARY", "SKILLS SUMMARY", "PAID EXPERIENCE", "RESEARCH & DEVELOPMENT", "EDUCATION"]:
            h = doc.add_paragraph(line)
            run = h.runs[0]
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            current_section = line
        elif line.startswith("•"):
            p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
        elif " - " in line and " | " in line: # Job title line
            p = doc.add_paragraph()
            p.add_run(line).bold = True
            p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph(line)
            if current_section == "PROFESSIONAL SUMMARY":
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(filename)
    print(f"Generated: {filename}")

def create_cover_letter_docx(cl_text, filename):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    for para in cl_text.split('\n\n'):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.space_after = Pt(12)

    doc.save(filename)
    print(f"Generated: {filename}")

# Role 1: Agentic AI
with open("tmp/prospects/nxtlevel-agentic-ai/resume.txt", "r") as f:
    resume_agentic = f.read()
create_resume_docx(resume_agentic, "tmp/prospects/nxtlevel-agentic-ai/resume.docx")

with open("tmp/prospects/nxtlevel-agentic-ai/cover_letter.txt", "r") as f:
    cl_agentic = f.read()
create_cover_letter_docx(cl_agentic, "tmp/prospects/nxtlevel-agentic-ai/cover_letter.docx")

# Role 2: GenAI Workflows
with open("tmp/prospects/nxtlevel-genai-workflows/resume.txt", "r") as f:
    resume_genai = f.read()
create_resume_docx(resume_genai, "tmp/prospects/nxtlevel-genai-workflows/resume.docx")

with open("tmp/prospects/nxtlevel-genai-workflows/cover_letter.txt", "r") as f:
    cl_genai = f.read()
create_cover_letter_docx(cl_genai, "tmp/prospects/nxtlevel-genai-workflows/cover_letter.docx")
