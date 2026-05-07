from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_doc(name, contact, summary, technical_core, experience, projects, education, filename, is_resume=True):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    n_para = doc.add_paragraph("STEPHEN JAEB")
    n_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = n_para.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    c_para = doc.add_paragraph(contact)
    c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_para.paragraph_format.space_after = Pt(0)
    
    l_para = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    l_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    l_para.paragraph_format.space_after = Pt(12)

    def add_header(text):
        h = doc.add_paragraph(text)
        run = h.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    if is_resume:
        add_header("PROFESSIONAL SUMMARY")
        summary_p = doc.add_paragraph(summary)
        summary_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_header("TECHNICAL CORE")
        for item in technical_core:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)

        add_header("PROFESSIONAL EXPERIENCE")
        for exp in experience:
            t_p = doc.add_paragraph()
            t_p.add_run(exp['title']).bold = True
            t_p.add_run(f"\t{exp['date']}").bold = True
            t_p.paragraph_format.space_after = Pt(4)
            for bullet in exp['bullets']:
                p = doc.add_paragraph(bullet, style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
        
        add_header("R&D AND SPECIAL PROJECTS")
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(proj['name']).bold = True
            doc.add_paragraph(proj['desc'], style='List Bullet')

        add_header("EDUCATION")
        doc.add_paragraph(education)

    doc.save(filename)
    print(f"Generated: {filename}")

# FreeWill Content
freewill_exp = [
    {
        "title": "Symetra - Cottage Grove, WI | Lead Backend Engineer / Architect",
        "date": "8/2025 – Present",
        "bullets": [
            "Architecting a high-scale integration hub for business-critical data, utilizing AI-native tools (Copilot/Codex) to maintain 100 percent test coverage and automate documentation.",
            "Engineering event-driven data pipelines on AWS Fargate with Polars and Pydantic, ensuring 100 percent payload correctness through forensic validation.",
            "Championing operational maturity by integrating end-to-end observability in Datadog and standardized QA/replay workflows."
        ]
    },
    {
        "title": "Veda Data Solutions - Madison, WI | Senior Software Engineer",
        "date": "3/2022 - 3/2025",
        "bullets": [
            "Achieved an 80x performance gain (7 days to 2 hours) for a distributed provider directory analysis solution by identifying and eliminating bottlenecks.",
            "Developed a Python-based distributed task engine using AWS Fargate and Redis, reducing operational costs by 40 percent.",
            "Optimized legacy systems using AI-assisted engineering practices, increasing test coverage from 45 percent to 100 percent."
        ]
    },
    {
        "title": "EatStreet - Madison, WI | Senior Software Engineer",
        "date": "5/2021 – 3/2022",
        "bullets": [
            "Architected a centralized POS Integration Hub using Java and Spring Boot to coordinate high-volume logistics and menu data."
        ]
    }
]

# Generate Resume
create_doc(
    "Lead Software Engineer, AI-Native",
    "Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com",
    "Lead Software Engineer and AI Strategist with over 20 years of experience in high-seniority architectural thinking and autonomous orchestration. Expert in building AI-native systems, implementing adversarial verification patterns, and driving technical direction through spec-driven development. Specialist in forensic system integrity and production-grade AI-native workflows.",
    ["AI-Native Engineering: Claude Code, Cursor, LLM-orchestrated SDLC, agentic workflows", "Architecture: Spec-driven development, adversarial verification, deterministic orchestration", "Backend: Python (FastAPI/Polars), AWS (Fargate/EventBridge), Docker, CI/CD"],
    freewill_exp,
    [{"name": "Data Platform Playbook", "desc": "Ongoing work designing autonomous multi-agent systems using LangGraph and Temporal with AI-native development workflows."}],
    "University of Houston - B.S. in Computer Engineering Technology (May 2000)",
    "tmp/custom/freewill/Stephen_Jaeb_Lead_AI_Native_FreeWill.docx"
)

# Generate Cover Letter
def create_cl():
    doc = Document()
    doc.add_paragraph("Stephen Jaeb\nCottage Grove, WI | sdjaeb@gmail.com\n" + datetime.now().strftime("%B %d, %Y") + "\n\nTo the FreeWill Engineering Team,")
    
    content = [
        "I am writing to express my enthusiasm for the Lead Software Engineer, AI-Native position at FreeWill. With over 20 years of engineering experience and a deep commitment to high-seniority architectural thinking in an AI-native world, I am eager to lead your team's pivot toward agentic orchestration and advanced AI workflows.",
        
        "In my current role as Lead Backend Engineer / Architect at Symetra, I have been a vocal champion for integrating AI-native tools like Copilot and Codex into the SDLC. I believe that being 'AI-Native' isn't just about using the tools, but about re-imagining our architectural patterns—such as adversarial verification and spec-driven development—to ensure that AI-augmented systems are trustworthy and resilient.",
        
        "My R&D work on the Data Platform Playbook explicitly prioritizes system orchestration using AI-native tools like Claude Code and Cursor. I am actively designing autonomous multi-agent systems using LangGraph and Temporal, exploring how we can build deterministic harnesses around non-deterministic models to deliver production-grade AI features.",
        
        "I am particularly drawn to FreeWill's mission of creating a more philanthropic world. I look forward to discussing how my experience in building resilient, AI-native platforms and my 'Forensic' approach to system integrity can help FreeWill achieve its goals."
    ]
    for p in content:
        para = doc.add_paragraph(p)
        para.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph("Sincerely,\n\nStephen Jaeb")
    doc.save("tmp/custom/freewill/Stephen_Jaeb_Cover_Letter_FreeWill.docx")
    print("Cover letter generated at tmp/custom/freewill/Stephen_Jaeb_Cover_Letter_FreeWill.docx")

create_cl()
