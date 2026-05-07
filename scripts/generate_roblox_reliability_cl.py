from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_roblox_cl():
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    name = doc.add_paragraph("Stephen Jaeb")
    run = name.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    
    contact = doc.add_paragraph("sdjaeb@gmail.com | (608) 852-9850 | Madison, WI")
    contact.paragraph_format.space_after = Pt(0)
    
    links = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    links.paragraph_format.space_after = Pt(24)

    date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_p.paragraph_format.space_after = Pt(18)

    recipient = doc.add_paragraph("To the Roblox AI/ML Engineering Team,")
    recipient.paragraph_format.space_after = Pt(12)

    content = [
        "As a Staff Architect with over two decades of experience building reliable, high-scale distributed systems and a deep background in game-tech R&D, I am writing to express my strong interest in the Principal AI/ML Engineer, Reliability position. I am eager to apply my expertise in operational maturity and deterministic orchestration to the challenge of scaling Roblox's AI-driven creator ecosystem.",
        
        "My career is defined by building the 'Deterministic Harness' required to turn non-deterministic models into production-grade systems. Most recently, as a Staff Architect at Symetra and Veda, I have specialized in building Integration Hubs and vectorized data engines that ensure 100 percent payload correctness and provenance at massive scale. I approach AI reliability through a systems engineering lens: ensuring that high-velocity AI features are as trustworthy and available as any core backend infrastructure.",
        
        "Why Roblox? I have long admired the platform's ability to scale high-fidelity experiences to millions of users globally. My experience at Great Wolf Resorts, where I led the launch of Unity3D-based mobile applications and immersive gaming experiences, gave me a first-hand look at the intersection of performance and player engagement. I am excited by your team's mission to ensure the reliability of AI tools that empower human creativity, and I am eager to bring my history of 80x performance gains to your platform.",
        
        "I bring a unique combination of 20 years of architectural maturity and a strategic focus on building robust, observable AI systems. I am eager to help Roblox architect the next generation of reliable, intelligent platforms.",
        
        "Thank you for your time and for continuing to push the boundaries of what is possible in the metaverse."
    ]

    for para_text in content:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    closing = doc.add_paragraph("Sincerely,")
    closing.paragraph_format.space_before = Pt(12)
    valediction = doc.add_paragraph("Stephen Jaeb")
    valediction.runs[0].bold = True

    doc.save("tmp/prospects/roblox-reliability/Stephen_Jaeb_Cover_Letter_Roblox_Reliability.docx")
    print("Cover letter generated at tmp/prospects/roblox-reliability/Stephen_Jaeb_Cover_Letter_Roblox_Reliability.docx")

if __name__ == "__main__":
    create_roblox_cl()
