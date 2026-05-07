from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_roblox_cl():
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    name = doc.add_paragraph("Stephen Jaeb")
    run = name.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    
    contact = doc.add_paragraph("sdjaeb@gmail.com | (608) 852-9850 | Madison, WI")
    contact.paragraph_format.space_after = Pt(0)
    
    links = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    links.paragraph_format.space_after = Pt(24)

    date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_p.paragraph_format.space_after = Pt(18)

    recipient = doc.add_paragraph("To the Roblox Creator AI Engineering Team,")
    recipient.paragraph_format.space_after = Pt(12)

    content = [
        "As a Staff Architect with over two decades of experience building high-scale distributed systems and a deep personal background in game-tech R&D (Unity3D), I am writing to express my strong interest in the Staff Engineer, Creator AI position. I am eager to apply my expertise in Python-based orchestration, agentic reasoning, and low-latency API design to Roblox's mission-critical creator ecosystem.",
        
        "My background is defined by the ability to build the 'Deterministic Harness' required to turn non-deterministic AI models into reliable production tools. Most recently, as a Staff Architect at Symetra and Veda, I have specialized in building Integration Hubs and vectorized task engines that ensure 100 percent data integrity and provenance at enterprise scale. I approach Creator AI through a systems engineering lens: optimizing the bridge between engine-level performance and AI-driven intelligence.",
        
        "Why Roblox? I have long admired Roblox's commitment to empowering creators through technical excellence. My experience at Great Wolf Resorts, where I led a cross-functional team in the launch of Unity3D-based mobile applications and immersive experiences, gave me a first-hand look at the unique challenges of scaling high-fidelity gaming infrastructure. I am excited by your team's focus on building AI tools that enhance human creativity, and I am eager to bring my history of 80x performance gains to your platform.",
        
        "I bring a unique combination of 20 years of architectural maturity and a strategic focus on implementing multi-agent orchestration (LangGraph). I am eager to help Roblox architect the next generation of intelligent tools for your global creator community.",
        
        "Thank you for your time and for continuing to redefine the future of play."
    ]

    for para_text in content:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    closing = doc.add_paragraph("Sincerely,")
    closing.paragraph_format.space_before = Pt(12)
    valediction = doc.add_paragraph("Stephen Jaeb")
    valediction.runs[0].bold = True

    doc.save("tmp/prospects/roblox/Stephen_Jaeb_Cover_Letter_Roblox.docx")
    print("Cover letter generated at tmp/prospects/roblox/Stephen_Jaeb_Cover_Letter_Roblox.docx")

if __name__ == "__main__":
    create_roblox_cl()
