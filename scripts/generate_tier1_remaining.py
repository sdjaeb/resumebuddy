from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc(name, contact, summary, technical_core, experience, projects, education, filename, is_resume=True):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    n_para = doc.add_paragraph("STEPHEN JAEB")
    n_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = n_para.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    c_para = doc.add_paragraph(contact)
    c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_para.paragraph_format.space_after = Pt(0)
    
    l_para = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    l_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    l_para.paragraph_format.space_after = Pt(12)

    def add_header(text):
        h = doc.add_paragraph(text)
        run = h.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    if is_resume:
        add_header("PROFESSIONAL SUMMARY")
        summary_p = doc.add_paragraph(summary)
        summary_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_header("TECHNICAL CORE")
        for item in technical_core:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)

        add_header("PROFESSIONAL EXPERIENCE")
        for exp in experience:
            t_p = doc.add_paragraph()
            t_p.add_run(exp['title']).bold = True
            t_p.add_run(f"\t{exp['date']}").bold = True
            t_p.paragraph_format.space_after = Pt(4)
            for bullet in exp['bullets']:
                p = doc.add_paragraph(bullet, style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
        
        add_header("PERSONAL PROJECTS")
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(proj['name']).bold = True
            doc.add_paragraph(proj['desc'], style='List Bullet')

        add_header("EDUCATION")
        doc.add_paragraph(education)

    doc.save(filename)
    print(f"Generated: {filename}")

# Shared Base Experience
base_exp = [
    {
        "title": "Symetra - Cottage Grove, WI | Lead Backend Engineer / Architect",
        "date": "8/2025 – Present",
        "bullets": [
            "Architected a high-scale Integration Hub using Python and FastAPI to move business-critical data between internal systems and external partners.",
            "Engineered event-driven data pipelines on AWS Fargate leveraging Polars and Pydantic for forensic schema validation and robust data normalization.",
            "Implemented 'Governance as Code' patterns by automating policy enforcement through Pydantic and real-time observability in Datadog."
        ]
    },
    {
        "title": "Veda Data Solutions - Madison, WI | Senior Software Engineer",
        "date": "3/2022 - 3/2025",
        "bullets": [
            "Achieved an 80x performance gain (7 days to 2 hours) by identifying bottlenecks with Pyinstrument and optimizing distributed processing.",
            "Developed an event-driven DDD framework in Python that unified complex business domains and reduced maintenance costs by 30 percent.",
            "Optimized legacy codebases using AI-assisted engineering, increasing test coverage from 45 percent to 100 percent."
        ]
    },
    {
        "title": "EatStreet - Madison, WI | Senior Software Engineer",
        "date": "5/2021 – 3/2022",
        "bullets": [
            "Architected scalable POS integration solutions (Centralized Integration Hub) using Java and Spring Boot to coordinate high-volume logistics data."
        ]
    }
]

# Nuclearn Kit
create_doc(
    "Staff AI Infrastructure Engineer",
    "Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com",
    "Staff Architect with over 20 years of experience delivering high-velocity AI infrastructure and distributed data engines. Expert in LLM serving, routing, and building the 'Deterministic Harness' for agentic orchestration. I specialize in bridging high-scale backend engineering (Python/AWS) with operational maturity to ensure AI models are safe, observable, and production-ready.",
    ["AI Infrastructure: LLM Serving/Routing, LangGraph, RAG, Prompt Engineering, Model Drift Monitoring", "Orchestration: Python (FastAPI/Polars/Pydantic), AWS (Fargate/EventBridge), Temporal, Redis", "Scale: Distributed state machines, 80x performance gain history, 100% test coverage"],
    base_exp,
    [{"name": "Data Platform Playbook", "desc": "Modular platform demonstrating multi-agent orchestration and high-throughput data engineering."}],
    "University of Houston - B.S. in Computer Engineering Technology (May 2000)",
    "tmp/prospects/nuclearn/Stephen_Jaeb_Staff_AI_Infra_Nuclearn.docx"
)

# Plaid Kit
create_doc(
    "Staff Software Engineer, Platform",
    "Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com",
    "Staff Architect with 20+ years of experience delivering mission-critical Fintech integration hubs and distributed data systems. Expert in architecting 'Zero-Failure' event-driven platforms that ensure 100 percent schema integrity and data provenance. I specialize in building the resilient orchestration layers required for high-compliance financial data flows on AWS.",
    ["Platform Engineering: High-scale Integration Hubs, Distributed Systems, Event-Driven Architecture", "Fintech Compliance: PCI-DSS, SOC2, Forensic Data Integrity, Pydantic Schema Enforcement", "Backend Stack: Python (FastAPI), AWS (Fargate/Kinesis/S3), Redis, 100% Test Coverage"],
    base_exp,
    [{"name": "Data Platform Playbook", "desc": "Modular reference architecture for high-scale, observable data pipelines."}],
    "University of Houston - B.S. in Computer Engineering Technology (May 2000)",
    "tmp/prospects/plaid/Stephen_Jaeb_Staff_Platform_Plaid.docx"
)

# Together AI Kit
create_doc(
    "Staff Systems Engineer",
    "Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com",
    "Staff Architect with over 20 years of experience delivering high-performance distributed systems and 'Scientific' data engines. Expert in GPU cluster orchestration, vectorized compute (Polars/Arrow), and building the architectural foundations for massive-scale ML training and inference. I specialize in achieving extreme performance gains (80x history) while maintaining 100 percent data fidelity.",
    ["Systems Engineering: Distributed Systems, GPU Cluster Orchestration, High-Performance Computing", "Data Integrity: Polars, Pydantic, Apache Arrow, Scientific Data Provenance, HIPAA/SOC2", "Cloud Infrastructure: AWS (Fargate/EventBridge), Temporal, Redis, Docker, CI/CD"],
    base_exp,
    [{"name": "Data Platform Playbook", "desc": "Modular platform featuring high-throughput vectorized data processing and orchestration."}],
    "University of Houston - B.S. in Computer Engineering Technology (May 2000)",
    "tmp/prospects/together_ai/Stephen_Jaeb_Staff_Systems_TogetherAI.docx"
)
