from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_resume():
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header
    name = doc.add_paragraph("STEPHEN JAEB")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    contact = doc.add_paragraph("Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(0)
    
    links = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_after = Pt(12)

    def add_section_header(text):
        h = doc.add_paragraph(text)
        run = h.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    # Professional Summary
    add_section_header("PROFESSIONAL SUMMARY")
    summary = doc.add_paragraph("Staff Architect & AI Strategist with over 20 years of experience delivering high-performance, compliant intelligent systems. Focused on multi-agent orchestration (LangGraph), distributed system design, and high-scale backend architecture. I specialize in building the 'Deterministic Harness' around AI models in regulated environments (FinTech, Insurance, HealthTech). Proven track record of leading technical strategy for 'Zero-Failure' integration hubs and high-velocity data engines on AWS.")
    summary.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Technical Core
    add_section_header("TECHNICAL CORE")
    core_items = [
        "AI Architecture: Multi-agent orchestration, LangGraph, RAG pipelines, LLM Safety/Observability",
        "Backend & Cloud: Python (FastAPI/Polars/Pydantic), AWS (Fargate/EventBridge/Redshift), SQL",
        "Compliance & Scale: PCI-DSS, SOC2, HIPAA, high-throughput integration hubs, 100% test coverage",
        "Leadership: Staff Architect, mentoring, technical roadmap definition, 'Zero-Failure' SDLC"
    ]
    for item in core_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # Experience
    add_section_header("PROFESSIONAL EXPERIENCE")

    # Symetra
    exp1 = doc.add_paragraph()
    exp1.add_run("Symetra - Cottage Grove, WI | Lead Backend Engineer / Architect").bold = True
    exp1.add_run("\t8/2025 – Present").bold = True
    exp1.paragraph_format.space_after = Pt(4)
    bullets1 = [
        "Architecting a high-scale Integration Hub for business-critical insurance and claims data, ensuring 100% payload correctness and implementing resilient fallbacks.",
        "Engineering scalable, event-driven data pipelines on AWS Fargate for robust data normalization and forensic schema validation.",
        "Implementing 'Governance as Code' by automating policy enforcement through Pydantic schemas and real-time observability in Datadog."
    ]
    for b in bullets1:
        doc.add_paragraph(b, style='List Bullet').paragraph_format.space_after = Pt(2)

    # Veda
    exp2 = doc.add_paragraph()
    exp2.add_run("Veda Data Solutions - Madison, WI | Senior Software Engineer").bold = True
    exp2.add_run("\t3/2022 - 3/2025").bold = True
    exp2.paragraph_format.space_before = Pt(8)
    exp2.paragraph_format.space_after = Pt(4)
    bullets2 = [
        "Architected a Python-based distributed task engine for provider directory analysis using AWS Fargate and Redis for state persistence.",
        "Achieved an 80x performance gain (7 days to 2 hours) by identifying bottlenecks with Pyinstrument and optimizing distributed processing.",
        "Ensured 100% data provenance for complex HealthTech datasets, satisfying rigorous accuracy requirements for model training and compliance."
    ]
    for b in bullets2:
        doc.add_paragraph(b, style='List Bullet').paragraph_format.space_after = Pt(2)

    # EatStreet
    exp3 = doc.add_paragraph()
    exp3.add_run("EatStreet - Madison, WI | Senior Software Engineer").bold = True
    exp3.add_run("\t5/2021 – 3/2022").bold = True
    exp3.paragraph_format.space_before = Pt(8)
    exp3.paragraph_format.space_after = Pt(4)
    doc.add_paragraph("Architected scalable POS integration solutions (Centralized Integration Hub) using Java and Spring Boot to ingest menu data and coordinate orders with external partners.", style='List Bullet')

    # JHT
    exp4 = doc.add_paragraph()
    exp4.add_run("Johnson Health Tech - Cottage Grove, WI | Data Analytics and Insights Manager").bold = True
    exp4.add_run("\t2/2019 – 4/2021").bold = True
    exp4.paragraph_format.space_before = Pt(8)
    exp4.paragraph_format.space_after = Pt(4)
    doc.add_paragraph("Established and led a dedicated Data Analytics and BI team, spearheading robust data architecture leveraging AWS Kinesis and Redshift.", style='List Bullet')

    # Projects
    add_section_header("PERSONAL PROJECTS")
    p_proj = doc.add_paragraph()
    p_proj.add_run("Data Platform Playbook | https://github.com/sdjaeb/data-platform-playbook").bold = True
    doc.add_paragraph("Modular data platform with modern practices: LangGraph, Temporal, Airflow, Kafka, FastAPI.", style='List Bullet')

    # Education
    add_section_header("EDUCATION")
    edu = doc.add_paragraph("University of Houston - B.S. in Computer Engineering Technology (May 2000)")

    doc.save("tmp/prospects/spoton/Stephen_Jaeb_Staff_AI_Engineer_SpotOn.docx")
    print("Resume generated at tmp/prospects/spoton/Stephen_Jaeb_Staff_AI_Engineer_SpotOn.docx")

if __name__ == "__main__":
    create_resume()
