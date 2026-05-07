from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc(title, contact_info, professional_summary, technical_core, experience_items, projects, education, filename, is_resume=True):
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header
    name = doc.add_paragraph("STEPHEN JAEB")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    contact = doc.add_paragraph(contact_info)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(0)
    
    links = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_after = Pt(12)

    def add_section_header(text):
        h = doc.add_paragraph(text)
        run = h.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    # Resume Specific Formatting
    if is_resume:
        add_section_header("PROFESSIONAL SUMMARY")
        summary = doc.add_paragraph(professional_summary)
        summary.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_section_header("TECHNICAL CORE")
        for item in technical_core:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)

        add_section_header("PROFESSIONAL EXPERIENCE")
        for exp in experience_items:
            title_p = doc.add_paragraph()
            title_p.add_run(exp['title']).bold = True
            title_p.add_run(f"\t{exp['date']}").bold = True
            title_p.paragraph_format.space_after = Pt(4)
            for bullet in exp['bullets']:
                p = doc.add_paragraph(bullet, style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
        
        add_section_header("PERSONAL PROJECTS")
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(proj['name']).bold = True
            doc.add_paragraph(proj['desc'], style='List Bullet')

        add_section_header("EDUCATION")
        doc.add_paragraph(education)

    doc.save(filename)
    print(f"Generated: {filename}")

# CORE DATA FROM resume.txt - Re-emphasized for specific tracks
# Applied AI: Focus on Orchestration, Determinism, and AI-Native SDLC.
applied_ai_experience = [
    {
        "title": "Symetra - Cottage Grove, WI | Lead Backend Engineer / Architect",
        "date": "8/2025 – Present",
        "bullets": [
            "Architected a high-scale Integration Hub using Python and FastAPI to move business-critical data between internal systems and external partners.",
            "Designed the core claims-payments pipeline, implementing resilient fallbacks and ensuring 100 percent payload correctness for high-stakes financial data.",
            "Engineered an event-driven data pipeline on AWS Fargate leveraging Polars and Pydantic for forensic schema validation and robust data normalization.",
            "Implemented Governance as Code patterns by automating policy enforcement through Pydantic and establishing real-time observability in Datadog.",
            "Accelerated delivery by integrating AI-native tools like Copilot and Codex into the SDLC, ensuring 100 percent test coverage and automated documentation."
        ]
    },
    {
        "title": "Veda Data Solutions - Madison, WI | Senior Software Engineer",
        "date": "3/2022 - 3/2025",
        "bullets": [
            "Developed a Python-based distributed task engine for provider directory analysis using AWS Fargate, EventBridge, and Redis for state persistence.",
            "Reduced processing time from 7 days to 2 hours (80x gain) and cut operational costs by 40 percent through optimization of distributed AI service integrations.",
            "Pioneered an event-driven, Domain-Driven Design (DDD) framework in Python that unified complex business domains and streamlined architecture by 30 percent.",
            "Optimized legacy codebases using AmazonQ, increasing test coverage from 45 percent to 100 percent while aggressively reducing technical debt."
        ]
    },
    {
        "title": "EatStreet - Madison, WI | Senior Software Engineer",
        "date": "5/2021 – 3/2022",
        "bullets": [
            "Architected a scalable POS Integration Hub using Java and Spring Boot to coordinate high-volume menu data and orders with external partners.",
            "Developed modular adapter services for third-party vendors, facilitating rapid company growth and seamless data exchange."
        ]
    },
    {
        "title": "Johnson Health Tech - Cottage Grove, WI | Data Analytics and Insights Manager",
        "date": "2/2019 – 4/2021",
        "bullets": [
            "Spearheaded data architecture development leveraging AWS Kinesis for real-time streaming and AWS Redshift for scalable warehousing.",
            "Partnered with the VP of Global R&D to establish a dedicated Data Analytics team, translating complex telemetry into actionable business insights."
        ]
    }
]

# Assets: Focus on Performance, Data Engineering, and Media Integration.
assets_experience = [
    {
        "title": "Symetra - Cottage Grove, WI | Lead Backend Engineer / Architect",
        "date": "8/2025 – Present",
        "bullets": [
            "Engineer scalable, event-driven data pipelines on AWS Fargate using Polars and Pydantic for high-velocity data normalization and asset ingestion.",
            "Design and implement high-throughput integration hubs (Python/FastAPI) to move business-critical data with 100 percent forensic integrity.",
            "Utilize Pyinstrument for continuous performance profiling, eliminating bottlenecks in mission-critical financial data flows.",
            "Lead the transition to a Zero-Failure architecture for claims payloads, persisting high-fidelity results to S3 and DynamoDB."
        ]
    },
    {
        "title": "Veda Data Solutions - Madison, WI | Senior Software Engineer",
        "date": "3/2022 - 3/2025",
        "bullets": [
            "Achieved an 80x performance gain (7 days to 2 hours) by refactoring distributed processing workflows using vectorized compute patterns (Pandas/Arrow).",
            "Built a distributed provider directory task engine using AWS Fargate, EventBridge, and S3, achieving a 40 percent reduction in operational costs.",
            "Ensured 100 percent data provenance for ML training sets by implementing comprehensive observability in Datadog and structured logging.",
            "Delivered a unified ingestion API (FastAPI) that abstracted away legacy database complexity into a standardized Data Product."
        ]
    },
    {
        "title": "EatStreet - Madison, WI | Senior Software Engineer",
        "date": "5/2021 – 3/2022",
        "bullets": [
            "Developed a high-volume POS Integration Hub in Java/Spring Boot to handle complex menu and image assets for external delivery partners.",
            "Optimized legacy integration workflows to ensure system stability and reduce downtime during peak logistics windows."
        ]
    },
    {
        "title": "Johnson Health Tech - Cottage Grove, WI | Data Analytics and Insights Manager",
        "date": "2/2019 – 4/2021",
        "bullets": [
            "Led organizational data initiatives leveraging AWS Kinesis (real-time streaming) and Redshift to process global equipment telemetry at scale.",
            "Translated millions of machine data points into predictive maintenance insights and high-fidelity QuickSight dashboards."
        ]
    }
]

# Generate Applied AI Files
create_doc(
    "Staff Software Engineer, Applied AI",
    "Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com",
    "Staff Architect with 20+ years of experience delivering high-velocity backend infrastructure. Expert in designing high-scale distributed systems and leveraging AI-native engineering tools (Cursor, Claude Code) to achieve 5x productivity gains. I specialize in building the 'Deterministic Harness' for agentic orchestration, enabling rapid feature delivery without compromising on schema-perfect validation or operational maturity.",
    ["AI Orchestration: LangChain, LangGraph (Stateful Agents), RAG, LLM Safety & Observability", "Backend Architecture: Python (FastAPI/Polars/Pydantic), AWS (Fargate/EventBridge), Redis", "Engineering Velocity: AI-native SDLC (Claude Code/Cursor), 100% Test Coverage, DDD"],
    applied_ai_experience,
    [{"name": "Data Platform Playbook", "desc": "Modular platform demonstrating multi-agent orchestration with LangGraph, Temporal, and FastAPI."}],
    "University of Houston - B.S. in Computer Engineering Technology (May 2000)",
    "tmp/prospects/webflow/applied_ai/Stephen_Jaeb_Staff_Applied_AI_Webflow.docx"
)

# Generate Assets Files
create_doc(
    "Staff Software Engineer, Assets",
    "Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com",
    "Staff Architect with over 20 years of experience delivering mission-critical data products. Expert in building high-scale, vectorized integration hubs (Python, Polars, Pydantic) and productionizing ML-driven asset capabilities. Proven history of delivering 80x performance gains and architecting distributed systems that ensure 100% data integrity and provenance at enterprise scale.",
    ["Data Engineering: Polars, Pydantic, High-velocity pipelines, S3/DynamoDB, Snowflake", "Distributed Systems: Performance profiling (Pyinstrument), AWS Fargate, Redis, Event-driven architecture", "Operational Excellence: 100% Test Coverage, Forensic Data Integrity, Datadog Observability"],
    assets_experience,
    [{"name": "Data Platform Playbook", "desc": "High-performance data platform featuring Polars, Kafka, and FastAPI for real-time asset processing."}],
    "University of Houston - B.S. in Computer Engineering Technology (May 2000)",
    "tmp/prospects/webflow/assets/Stephen_Jaeb_Staff_Assets_Webflow.docx"
)
