from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_cover_letter():
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Header
    name = doc.add_paragraph("Stephen Jaeb")
    run = name.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    
    contact = doc.add_paragraph("sdjaeb@gmail.com | (608) 852-9850 | Madison, WI")
    contact.paragraph_format.space_after = Pt(0)
    
    links = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    links.paragraph_format.space_after = Pt(24)

    # Date
    from datetime import datetime
    date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_p.paragraph_format.space_after = Pt(18)

    # Recipient
    recipient = doc.add_paragraph("To the SpotOn Engineering Team,")
    recipient.paragraph_format.space_after = Pt(12)

    # Content - splitting by double newlines for paragraphs
    content = [
        "As a Staff Architect with over two decades of experience building mission-critical backend systems and a specialized focus on multi-agent orchestration, I am writing to express my strong interest in the Staff AI Engineer position. I am eager to apply my expertise in high-scale distributed systems and Python-based AI infrastructure to SpotOn's high-growth 'Accelerate' team.",
        
        "My background is defined by the ability to build robust, compliant 'Deterministic Harnesses' around the non-deterministic nature of modern AI models. Most recently, as a Staff Architect in the Insurance and HealthTech sectors (Symetra and Veda), I have specialized in building Integration Hubs and vectorized data engines that ensure 100% schema integrity and data provenance. I treat AI not as a black box, but as an orchestration challenge—leveraging patterns like LangGraph and Temporal to build systems that are as trustworthy as they are intelligent.",
        
        "Why SpotOn? I am impressed by your team's mission to redefine engineering velocity through AI-native practices. I approach the SDLC through the same lens: I don't just use AI to write code; I use multi-agent systems to automate the 'Hard Parts' of engineering—adversarial verification, drift detection, and forensic observability. My experience in high-compliance environments (PCI-DSS and HIPAA) aligns perfectly with your need for an AI-native strategist who understands that speed is useless without safety and reliability.",
        
        "I bring a unique combination of 20 years of architectural discipline and a strategic focus on implementing multi-agent orchestration (LangGraph). I am eager to help SpotOn build the intelligent platforms that fuel your next stage of scale.",
        
        "Thank you for your time and for leading the charge in the future of AI-driven Fintech."
    ]

    for para_text in content:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Closing
    closing = doc.add_paragraph("Sincerely,")
    closing.paragraph_format.space_before = Pt(12)
    
    valediction = doc.add_paragraph("Stephen Jaeb")
    valediction.runs[0].bold = True

    doc.save("tmp/prospects/spoton/Stephen_Jaeb_Cover_Letter_SpotOn.docx")
    print("Cover letter generated at tmp/prospects/spoton/Stephen_Jaeb_Cover_Letter_SpotOn.docx")

if __name__ == "__main__":
    create_cover_letter()
