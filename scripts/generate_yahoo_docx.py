from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc(name, contact, summary, technical_core, experience, projects, education, filename, is_resume=True):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    n_para = doc.add_paragraph("STEPHEN JAEB")
    n_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = n_para.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    c_para = doc.add_paragraph(contact)
    c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_para.paragraph_format.space_after = Pt(0)
    
    l_para = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    l_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    l_para.paragraph_format.space_after = Pt(12)

    def add_header(text):
        h = doc.add_paragraph(text)
        run = h.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    if is_resume:
        add_header("PROFESSIONAL SUMMARY")
        summary_p = doc.add_paragraph(summary)
        summary_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_header("TECHNICAL CORE")
        for item in technical_core:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)

        add_header("PROFESSIONAL EXPERIENCE")
        for exp in experience:
            t_p = doc.add_paragraph()
            t_p.add_run(exp['title']).bold = True
            t_p.add_run(f"\t{exp['date']}").bold = True
            t_p.paragraph_format.space_after = Pt(4)
            for bullet in exp['bullets']:
                p = doc.add_paragraph(bullet, style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
        
        add_header("PERSONAL PROJECTS")
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(proj['name']).bold = True
            doc.add_paragraph(proj['desc'], style='List Bullet')

        add_header("EDUCATION")
        doc.add_paragraph(education)

    doc.save(filename)
    print(f"Generated: {filename}")

# COMPREHENSIVE YAHOO EXPERIENCE (Full History from resume.txt)
yahoo_exp_full = [
    {
        "title": "Symetra - Cottage Grove, WI | Lead Backend Engineer / Architect",
        "date": "8/2025 – Present",
        "bullets": [
            "Lead and architect the development of an enterprise-scale backend integration hub using Python and FastAPI to synchronize business-critical data between internal systems and external partners.",
            "Design and implement core claims-payments pipelines, ensuring 100 percent payload correctness and implementing resilient fallbacks for high-stakes financial data.",
            "Engineer scalable, event-driven data pipelines on AWS Fargate leveraging Polars and Pydantic for robust data normalization and forensic schema validation.",
            "Champion operational maturity by integrating end-to-end observability with Datadog and managing CI/CD infrastructure via CloudFormation.",
            "Drive technical excellence by integrating AI tools into the SDLC and maintaining 100 percent test coverage for all backend integrations."
        ]
    },
    {
        "title": "Veda Data Solutions - Madison, WI | Senior Software Engineer",
        "date": "3/2022 - 3/2025",
        "bullets": [
            "Developed a Python-based distributed task engine for a provider directory analysis solution using AWS Fargate, EventBridge, S3, and Redis.",
            "Achieved an 80x performance gain (7 days to 2 hours) and reduced operational costs by 40 percent through refactoring distributed processing workflows.",
            "Pioneered an event-driven DDD framework in Python that unified complex business domains and reduced long-term maintenance costs by 30 percent.",
            "Increased test coverage from 45 percent to 100 percent by leveraging AI-assisted engineering tools and aggressively reducing technical debt.",
            "Ensured data integrity for ML model training by implementing structured logging and comprehensive observability suites."
        ]
    },
    {
        "title": "EatStreet - Madison, WI | Senior Software Engineer",
        "date": "5/2021 – 3/2022",
        "bullets": [
            "Architected scalable POS integration solutions handlers handled high-volume menu data and orders across a Service-Oriented Architecture using Java and Spring Boot.",
            "Developed modular adapter services for third-party vendors, facilitating rapid company growth and seamless data exchange.",
            "Resolved complex issues and applied critical fixes in legacy integration systems, optimizing workflows to ensure system stability."
        ]
    },
    {
        "title": "Johnson Health Tech - Cottage Grove, WI | Data Analytics and Insights Manager",
        "date": "2/2019 – 4/2021",
        "bullets": [
            "Partnered with the VP of Global R&D to establish and lead a dedicated Data Analytics and BI team, defining team structure and talent needs.",
            "Spearheaded robust data architecture leveraging AWS Kinesis (real-time streaming) and Redshift (scalable warehousing) to translate telemetry into actionable product insights.",
            "Led organizational data initiatives, translating complex data into insights for senior management via QuickSight dashboards and ELK Stack log analysis."
        ]
    },
    {
        "title": "Johnson Health Tech - Cottage Grove, WI | Senior Web Engineer",
        "date": "3/2016 – 2/2019",
        "bullets": [
            "Architected an internal platform for global software updates using Node.js and AWS API Gateway.",
            "Engineered scalable ETL pipelines using AWS Lambda and Kinesis to process high-volume machine telemetry from global sources."
        ]
    },
    {
        "title": "Nextpoint - Madison, WI | Software Engineer",
        "date": "9/2013 – 2/2016",
        "bullets": [
            "Maintained and extended a Ruby on Rails SaaS platform for eDiscovery, managing terabyte-scale storage across S3 buckets for compliance and scalability."
        ]
    },
    {
        "title": "Great Wolf Resorts - Madison, WI | Web Developer III",
        "date": "4/2010 – 9/2013",
        "bullets": [
            "Led full-stack development and architecture optimization for global web platforms, implementing HAProxy and Varnish to accelerate content delivery.",
            "Developed custom backend modules in Drupal (PHP) and integrated with Java/Spring REST services for high-availability booking systems.",
            "Spearheaded R&D for Unity3D-based mobile applications and on-site immersive experiences."
        ]
    }
]

# Generate Full Yahoo Resume
create_doc(
    "Senior Data Engineer (AI-Driven Automation)",
    "Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com",
    "Staff Architect and Data Strategist with over 20 years of experience delivering high-scale distributed systems and intelligent data architecture. Expert in designing high-velocity data pipelines (Python/AWS) and implementing AI-driven automation that ensures forensic data integrity and 100 percent payload correctness. Proven track record of achieving 80x performance gains and leading technical direction in regulated, mission-critical environments.",
    ["Data Engineering: Python (FastAPI/Polars/Pydantic), SQL, High-throughput automation, GCP/AWS", "Distributed Systems: Event-Driven Architecture, Redis, Temporal, AWS Fargate/Redshift", "Operational Maturity: 100% Test Coverage, Forensic Data Integrity, Datadog Observability", "Architecture & Strategy: DDD, Hexagonal, Microservices, Technical Mentorship"],
    yahoo_exp_full,
    [{"name": "Data Platform Playbook", "desc": "Modular platform implemented with modern practices featuring Airflow, Kafka, Spark, dbt, and FastAPI."}],
    "University of Houston - B.S. in Computer Engineering Technology (May 2000)",
    "tmp/prospects/yahoo/Stephen_Jaeb_Senior_Data_Engineer_Yahoo.docx"
)
