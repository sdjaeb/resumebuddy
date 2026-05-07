from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc(title, contact_info, professional_summary, technical_core, experience_items, projects, education, filename, is_resume=True):
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header
    name = doc.add_paragraph("STEPHEN JAEB")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    contact = doc.add_paragraph(contact_info)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(0)
    
    links = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_after = Pt(12)

    def add_section_header(text):
        h = doc.add_paragraph(text)
        run = h.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    if is_resume:
        add_section_header("PROFESSIONAL SUMMARY")
        summary = doc.add_paragraph(professional_summary)
        summary.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_section_header("TECHNICAL CORE")
        for item in technical_core:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)

        add_section_header("PROFESSIONAL EXPERIENCE")
        for exp in experience_items:
            title_p = doc.add_paragraph()
            title_p.add_run(exp['title']).bold = True
            title_p.add_run(f"\t{exp['date']}").bold = True
            title_p.paragraph_format.space_after = Pt(4)
            for bullet in exp['bullets']:
                p = doc.add_paragraph(bullet, style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
        
        add_section_header("PERSONAL PROJECTS")
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(proj['name']).bold = True
            doc.add_paragraph(proj['desc'], style='List Bullet')

        add_section_header("EDUCATION")
        doc.add_paragraph(education)

    else:
        doc.add_paragraph(f"\n{professional_summary}")

    doc.save(filename)
    print(f"Generated: {filename}")

# Roblox Content
roblox_experience = [
    {
        "title": "Symetra - Cottage Grove, WI | Lead Backend Engineer / Architect",
        "date": "8/2025 – Present",
        "bullets": [
            "Architecting high-scale Integration Hubs and agentic orchestration platforms on AWS Fargate using Python, FastAPI, and Pydantic.",
            "Designing deterministic harnesses for non-deterministic AI systems, ensuring 100 percent data integrity and governance.",
            "Accelerating engineering velocity by integrating AI-native tools (Cursor/Claude Code) into the architectural strategy."
        ]
    },
    {
        "title": "Veda Data Solutions - Madison, WI | Senior Software Engineer",
        "date": "3/2022 - 3/2025",
        "bullets": [
            "Achieved an 80x performance gain (7 days to 2 hours) by identifying bottlenecks with Pyinstrument and refactoring distributed processing workflows.",
            "Built a distributed task engine using AWS Fargate and Redis for state persistence, managing high-volume data ingestion for ML models.",
            "Developed an event-driven DDD framework that unified complex business domains and reduced maintenance costs by 30 percent."
        ]
    },
    {
        "title": "Great Wolf Resorts - Madison, WI | Web Developer III (Unity3D Lead)",
        "date": "4/2010 - 9/2013",
        "bullets": [
            "Led a cross-functional team in the R&D and launch of immersive Unity3D-based mobile applications and on-site gaming experiences.",
            "Architected high-availability web platforms using HAProxy and Varnish, optimizing content delivery for millions of guests."
        ]
    },
    {
        "title": "EatStreet - Madison, WI | Senior Software Engineer",
        "date": "5/2021 – 3/2022",
        "bullets": [
            "Architected scalable POS integration solutions handling high-volume logistics and asset data across a Service-Oriented Architecture."
        ]
    }
]

# Generate Roblox Files
create_doc(
    "Staff Engineer, Creator AI",
    "Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com",
    "Staff Architect with 20+ years of experience delivering high-scale distributed systems and intelligent creator tools. Expert in agentic orchestration (LangGraph), low-latency API serving, and bridging the gap between engine-level performance (Unity3D) and modern cloud architecture. Proven track record of leading technical strategy for mission-critical platforms in gaming, HealthTech, and FinTech.",
    ["AI & Agentic Systems: LangGraph, RAG pipelines, LLM Safety & Deterministic Orchestration", "Cloud Architecture: AWS (Fargate/EventBridge), Python (FastAPI/Polars/Pydantic), Redis", "Gaming & Performance: Unity3D/C#, performance profiling (Pyinstrument), bottleneck elimination"],
    roblox_experience,
    [{"name": "Data Platform Playbook", "desc": "Modular platform demonstrating multi-agent orchestration and high-throughput data engineering."}],
    "University of Houston - B.S. in Computer Engineering Technology (May 2000)",
    "tmp/prospects/roblox/Stephen_Jaeb_Staff_Creator_AI_Roblox.docx"
)
