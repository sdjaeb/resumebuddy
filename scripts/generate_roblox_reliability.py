from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_doc(name, contact, summary, technical_core, experience, projects, education, filename, is_resume=True):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    n_para = doc.add_paragraph("STEPHEN JAEB")
    n_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = n_para.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    c_para = doc.add_paragraph(contact)
    c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_para.paragraph_format.space_after = Pt(0)
    
    l_para = doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/stephen-jaeb/ | GitHub: https://github.com/sdjaeb/")
    l_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    l_para.paragraph_format.space_after = Pt(12)

    def add_header(text):
        h = doc.add_paragraph(text)
        run = h.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    if is_resume:
        add_header("PROFESSIONAL SUMMARY")
        summary_p = doc.add_paragraph(summary)
        summary_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_header("TECHNICAL CORE")
        for item in technical_core:
            p = doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)

        add_header("PROFESSIONAL EXPERIENCE")
        for exp in experience:
            t_p = doc.add_paragraph()
            t_p.add_run(exp['title']).bold = True
            t_p.add_run(f"\t{exp['date']}").bold = True
            t_p.paragraph_format.space_after = Pt(4)
            for bullet in exp['bullets']:
                p = doc.add_paragraph(bullet, style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
        
        add_header("R&D (PLAYBOOK) WORK")
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(proj['name']).bold = True
            doc.add_paragraph(proj['desc'], style='List Bullet')

        add_header("EDUCATION")
        doc.add_paragraph(education)

    doc.save(filename)
    print(f"Generated: {filename}")

# Roblox Reliability Content
roblox_exp = [
    {
        "title": "Symetra - Cottage Grove, WI | Lead Backend Engineer / Architect",
        "date": "8/2025 – Present",
        "bullets": [
            "Architecting high-scale Integration Hubs (Python/FastAPI) to ensure extreme data fidelity and near-zero-latency validation for insurance data flows.",
            "Designing deterministic harnesses for non-deterministic systems, implementing schema-perfect validation with Pydantic and forensic auditability.",
            "Championing operational maturity by integrating end-to-end observability in Datadog and identifying performance bottlenecks with Pyinstrument."
        ]
    },
    {
        "title": "Veda Data Solutions - Madison, WI | Senior Software Engineer",
        "date": "3/2022 - 3/2025",
        "bullets": [
            "Achieved an 80x performance gain (7 days to 2 hours) by refactoring legacy data processes into a vectorized distributed engine on AWS Fargate.",
            "Built a distributed task engine using Redis for state persistence, managing high-volume data ingestion for ML model training with 100 percent provenance.",
            "Developed an event-driven DDD framework that unified complex business domains and reduced long-term maintenance costs by 30 percent."
        ]
    },
    {
        "title": "Great Wolf Resorts - Madison, WI | Web Developer III (Unity3D Lead)",
        "date": "4/2010 - 9/2013",
        "bullets": [
            "Led a cross-functional team in the R&D and launch of immersive Unity3D-based mobile applications and on-site gaming experiences.",
            "Optimized content delivery and security for millions of guests by implementing HAProxy and Varnish for high-availability systems."
        ]
    }
]

create_doc(
    "Principal AI/ML Engineer, Reliability",
    "Cottage Grove, WI | (608) 852-9850 | sdjaeb@gmail.com",
    "Staff Architect with 20+ years of experience specialized in building reliable, high-scale distributed systems and intelligent data architecture. Expert in creating 'Deterministic Harnesses' for AI models and achieving extreme performance gains (80x history) in production environments. I specialize in bridging high-fidelity engineering (Python/AWS) with the performance and scalability requirements of modern AI platforms.",
    ["Reliability Engineering: High-Availability Distributed Systems, Fault-Tolerance, Observability (Datadog)", "AI Orchestration: LangGraph, RAG, LLM Safety, Deterministic Harnesses, Performance Tuning", "Backend Stack: Python (FastAPI/Polars/Pydantic), AWS (Fargate/EventBridge), Redis, Unity3D"],
    roblox_exp,
    [{"name": "Data Platform Playbook", "desc": "Modular platform demonstrating multi-agent orchestration and high-throughput data engineering."}],
    "University of Houston - B.S. in Computer Engineering Technology (May 2000)",
    "tmp/prospects/roblox-reliability/Stephen_Jaeb_Principal_Reliability_Roblox.docx"
)
