from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc(dir_path, filename_prefix):
    # Resume
    res_path = os.path.join(dir_path, "resume.txt")
    if os.path.exists(res_path):
        with open(res_path, 'r') as f:
            lines = f.readlines()
        
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Name
        name_p = doc.add_paragraph(lines[0].strip())
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_p.runs[0]
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Arial'

        # Contact
        contact_p = doc.add_paragraph(lines[1].strip())
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(0)
        
        links_p = doc.add_paragraph(lines[2].strip())
        links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_p.paragraph_format.space_after = Pt(12)

        for line in lines[3:]:
            l = line.strip()
            if not l:
                continue
            if l.isupper() and len(l) > 3:
                h = doc.add_paragraph(l)
                run = h.runs[0]
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Arial'
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)
            elif l.startswith("- "):
                p = doc.add_paragraph(l[2:], style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
            elif "|" in l and len(l) < 100:
                p = doc.add_paragraph()
                p.add_run(l).bold = True
                p.paragraph_format.space_after = Pt(4)
            else:
                p = doc.add_paragraph(l)
                p.paragraph_format.space_after = Pt(6)

        out_path = os.path.join(dir_path, f"Stephen_Jaeb_Resume_{filename_prefix}.docx")
        doc.save(out_path)
        print(f"Generated: {out_path}")

    # Cover Letter
    cl_path = os.path.join(dir_path, "cover_letter.txt")
    if os.path.exists(cl_path):
        with open(cl_path, 'r') as f:
            text = f.read()
        
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        for para in text.split('\n\n'):
            if not para.strip():
                continue
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(12)
            if len(para.strip()) > 100 and para.strip().endswith('.'):
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        out_path = os.path.join(dir_path, f"Stephen_Jaeb_Cover_Letter_{filename_prefix}.docx")
        doc.save(out_path)
        print(f"Generated: {out_path}")

prospects = [
    ("tmp/prospects/archetype-ai", "Archetype_AI"),
    ("tmp/prospects/global-investment-firm", "Global_Investment_Firm"),
    ("tmp/prospects/cresta", "Cresta"),
    ("tmp/prospects/ai-search-platform", "AI_Search_Platform")
]

for d, p in prospects:
    create_doc(d, p)
